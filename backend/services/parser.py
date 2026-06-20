import os
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import pytesseract

def parse_file(file_path: str, file_type: str) -> str:
    path = Path(file_path)

    if file_type == "pdf":
        return _parse_pdf(path)
    elif file_type == "docx":
        return _parse_docx(path)
    elif file_type in ("xlsx", "xls"):
        return _parse_excel(path, file_type)
    elif file_type in ("png", "jpg", "jpeg", "webp"):
        return _parse_image(path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def _parse_pdf(path: Path) -> str:
    text = []
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                content = page.extract_text()
                if content:
                    text.append(f"[Page {i+1}]\n{content}")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")
    if not text:
        raise ValueError("No text could be extracted from the PDF. The file may be scanned or encrypted.")
    return "\n\n".join(text)

def _parse_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")
    if not text.strip():
        raise ValueError("No text found in the DOCX file.")
    return text

def _parse_excel(path: Path, file_type: str) -> str:
    try:
        # openpyxl handles .xlsx; for .xls try xlrd if available, else load via openpyxl
        if file_type == "xls":
            try:
                import xlrd
                wb = xlrd.open_workbook(str(path))
                text = []
                for sheet_idx in range(wb.nsheets):
                    ws = wb.sheet_by_index(sheet_idx)
                    text.append(f"[Sheet: {ws.name}]")
                    for row_idx in range(ws.nrows):
                        row_text = "\t".join(str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols))
                        if row_text.strip():
                            text.append(row_text)
                return "\n".join(text)
            except ImportError:
                pass  # Fall through to openpyxl

        wb = openpyxl.load_workbook(str(path), data_only=True)
        text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text.append(f"[Sheet: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(c) if c is not None else "" for c in row)
                if row_text.strip():
                    text.append(row_text)
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Failed to parse Excel file: {e}")

def _parse_image(path: Path) -> str:
    try:
        img = Image.open(str(path))
        # Convert to RGB if needed (e.g. RGBA PNGs)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        # Don't hard-fail — return empty so upload.py can insert placeholder text
        return ""